from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/tiffanywright/Documents/New project/confident-clinician/Clinical Confidence Lab - Prompt Example Library.docx"
FOREST = RGBColor(42, 60, 47)
TERRA = RGBColor(197, 106, 77)
OLIVE = RGBColor(100, 115, 91)
SAND = "F1EADF"

examples = [
    ("The fantasy version of a confident therapist", "Calm, articulate, certain, never awkward, and always able to say something meaningful right away."),
    ("The trustworthy version I want to become", "I want to notice when I’m activated, slow down enough to think, and respond honestly instead of performing."),
    ("A confident clinician is not someone who…", "Never feels uncertain, always knows the perfect intervention, or makes every session feel profound."),
    ("A confident clinician is someone who…", "Can pause, stay curious, make a thoughtful choice, repair when needed, and use support."),
    ("The moment I feel least confident", "When a client becomes quiet and I cannot tell whether they need space, a question, or a change in direction."),
    ("When that happens, I tend to…", "I ask too many questions, explain more than necessary, and replay the session afterward."),
    ("The clinical skill I most want to strengthen", "Helping a client move from describing events into exploring the feeling or meaning underneath."),
    ("The confidence behavior I want to practice", "Pause for one breath and reflect what I heard before asking another question."),
    ("Evidence I want after four weeks", "I can notice my urge to perform and choose one grounded next move instead."),
    ("Knowledge or skill need", "I know the concept of emotional reflection, but I need clearer language for using it naturally in the room."),
    ("Practice or repetition need", "I have tried collaborative focusing only a few times, so it disappears when I feel pressured."),
    ("Clinical judgment need", "I can identify several possible directions but struggle to choose which matters most right now."),
    ("Confidence or activation need", "I know what I could do, but self-consciousness makes me rush and lose access to it."),
    ("Role or context need", "Back-to-back sessions and unclear expectations are making it harder to think and close on time."),
    ("Sustainability need", "I need a real transition after work instead of reviewing sessions throughout the evening."),
    ("The gap I am tempted to shame", "Needing more time to understand the thread when a client brings several concerns at once."),
    ("A more accurate description", "I am still building prioritization through repetition; this is a practice need, not proof that I am incapable."),
    ("A responsible practice step", "Use CLEAR on one de-identified example and bring my reasoning to supervision."),
    ("Activation pattern 1", "Overexplaining — trigger: silence; what I do: teach too quickly; cost: I stop listening for the client’s meaning."),
    ("Activation pattern 2", "Disappearing — trigger: possible disagreement; what I do: avoid redirecting; cost: the session loses direction."),
    ("The helpful-looking anxiety pattern", "Rescuing, because extra reassurance looks caring even when I am trying to make my own anxiety stop."),
    ("The earliest body cue", "My jaw tightens and I lean forward as if I need to make something happen immediately."),
    ("The pause I can practice", "Put both feet on the floor, exhale, and silently ask, “What do I actually know?”"),
    ("What happened before confidence dropped", "The client answered two questions with “I don’t know” and looked toward the floor."),
    ("Known or observable facts", "The client became quiet after we shifted topics; I noticed a longer pause and shorter answers."),
    ("The story or assumption my brain created", "My brain said, “They think I’m not helping,” even though I did not have evidence of that."),
    ("Body cues and activation", "My chest tightened, I spoke faster, and I felt an urge to fill the silence."),
    ("Protective patterns", "I moved into overexplaining so I could feel useful and avoid uncertainty."),
    ("Skill or training gaps", "I understand the idea in theory, but I need more repetition using it under pressure."),
    ("Support or consultation", "I will bring one focused question to supervision instead of researching alone after work."),
    ("Defining a confident clinician", "Someone who can pause, stay curious, repair when needed, and ask for support."),
    ("Behavior experiment", "Pause for one full breath before asking another question."),
    ("Clinical thread", "The details differ, but the repeating thread may be fear of disappointing other people."),
    ("Grounded response or next move", "Reflect what I notice, check whether it fits, and let the client correct me."),
    ("Priority", "Connection and clarification come first; advice can wait until I understand the need."),
    ("Role reminder", "My role is to assess, collaborate, document appropriately, and consult—not control the outcome."),
    ("What I can release", "I can release the demand to know exactly how the client experienced every moment."),
    ("What I might try differently", "I will slow down, summarize the two threads I hear, and ask where we should focus."),
    ("What happened", "The session moved across several topics, and I noticed urgency when we had ten minutes left."),
    ("Closing or ending language", "We have about ten minutes left; what feels most important to carry forward today?"),
    ("Clinical value in behavior", "Steadiness: I pause and respond thoughtfully; it does not require me to feel perfectly calm."),
    ("Boundary", "I will begin landing at the ten-minute mark and finish at the scheduled time."),
    ("Evidence of growth", "I noticed the urge to rescue, paused, and chose one collaborative question instead."),
    ("30-day or weekly commitment", "Once each week, I will use CLEAR on one de-identified moment and bring one question to consultation."),
    ("Open-ended reflection", "A brief, specific response from one fictional or fully de-identified moment—enough to show your thinking, not an entire case history."),
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.75)
sec.bottom_margin = Inches(.75)
sec.left_margin = Inches(.85)
sec.right_margin = Inches(.85)
sec.header_distance = Inches(.35)
sec.footer_distance = Inches(.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = FOREST
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [
    ("Heading 1", 18, FOREST, 16, 7),
    ("Heading 2", 14, TERRA, 12, 5),
    ("Heading 3", 11, FOREST, 8, 3),
]:
    st = styles[name]
    st.font.name = "Aptos Display"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
header.text = "THE CONFIDENT CLINICIAN  ·  FACILITATOR REFERENCE"
header.style = normal
header.runs[0].font.size = Pt(8)
header.runs[0].font.bold = True
header.runs[0].font.color.rgb = OLIVE

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(5)
run = title.add_run("Clinical Confidence Lab")
run.font.name = "Aptos Display"
run.font.size = Pt(27)
run.font.bold = True
run.font.color.rgb = FOREST
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
r = subtitle.add_run("Prompt Example Library")
r.font.name = "Aptos Display"
r.font.size = Pt(18)
r.font.color.rgb = TERRA

p = doc.add_paragraph()
p.add_run("Purpose. ").bold = True
p.add_run("These samples model the appropriate level of specificity for participant responses. They are examples—not preferred answers, clinical directives, or a substitute for supervision.")

note = doc.add_paragraph()
note.paragraph_format.left_indent = Inches(.18)
note.paragraph_format.right_indent = Inches(.18)
note.paragraph_format.space_before = Pt(7)
note.paragraph_format.space_after = Pt(14)
shade = OxmlElement("w:shd")
shade.set(qn("w:fill"), SAND)
note._p.get_or_add_pPr().append(shade)
rr = note.add_run("Privacy reminder: ")
rr.bold = True
rr.font.color.rgb = TERRA
note.add_run("Use fictional, composite, or fully de-identified material. Never include names, exact dates, workplaces, screenshots, records, or combinations of details that could identify a client.")

doc.add_heading("Course-wide examples", level=1)
doc.add_paragraph("The course app matches each interactive prompt to the most relevant example below. If a prompt is broad, the final open-ended example is used.")

for idx, (prompt_type, sample) in enumerate(examples, 1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    n = p.add_run(f"{idx:02d}  {prompt_type}")
    n.bold = True
    n.font.color.rgb = FOREST
    ex = doc.add_paragraph()
    ex.paragraph_format.left_indent = Inches(.28)
    ex.paragraph_format.space_after = Pt(5)
    label = ex.add_run("Example: ")
    label.bold = True
    label.font.color.rgb = TERRA
    sample_run = ex.add_run(sample)
    sample_run.italic = True
    sample_run.font.color.rgb = OLIVE

doc.add_heading("How to use or revise these", level=1)
for text in [
    "Keep examples short enough that participants still do their own thinking.",
    "Model process and clinical reasoning—not a single “correct” intervention.",
    "Use ordinary, relatable language rather than polished therapy-speak.",
    "Keep every clinical illustration fictional, composite, or fully de-identified.",
    "When a prompt involves risk, ethics, law, scope, competence, or policy, direct participants to formal supervision and required procedures.",
]:
    doc.add_paragraph(text, style="List Bullet")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("The Confident Clinician · Clinical Confidence Lab · Prompt Example Library").font.size = Pt(8)

doc.save(OUT)
print(OUT)
